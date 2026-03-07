import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False


def _safe(text) -> str:
    """Replace Unicode characters that fpdf2 standard fonts can't encode."""
    if not isinstance(text, str):
        text = str(text)
    # Replace common symbols with ASCII equivalents
    replacements = {
        '₹': 'Rs.',
        '→': '->',
        '←': '<-',
        '–': '-',
        '—': '--',
        '•': '*',
        '\u2019': "'",
        '\u2018': "'",
        '\u201c': '"',
        '\u201d': '"',
        '…': '...',
        '✅': '[OK]',
        '⚠': '[!]',
        '🔴': '[HIGH]',
        '🟡': '[MED]',
        '🟢': '[LOW]',
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    # Drop remaining non-Latin-1 chars
    text = text.encode('latin-1', errors='ignore').decode('latin-1')
    return text


class CAMReportGenerator:

    def __init__(self, output_dir: str = "reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, company_name: str, loan_amount: float,
                 ratios: dict, risk_result: dict, loan_decision: dict,
                 fraud_result: dict = None, circular_result: dict = None,
                 case_id: str = None):

        if case_id is None:
            case_id = f"CAM-{datetime.now().strftime('%Y-%m%d-%H%M')}"

        report_data = {
            "case_id": case_id,
            "generated_at": datetime.now().isoformat(),
            "company": company_name,
            "loan_amount_cr": loan_amount,
            "financial_ratios": ratios,
            "risk_assessment": risk_result,
            "loan_decision": loan_decision,
            "fraud_analysis": fraud_result or {},
            "circular_trading": circular_result or {},
        }

        docx_path = None
        json_path = None
        pdf_path = None
        try:
            docx_path = self._generate_docx(report_data)
        except Exception as e:
            print(f"DOCX generation error: {e}")
        try:
            json_path = self._generate_json(report_data)
        except Exception as e:
            print(f"JSON generation error: {e}")
        if HAS_FPDF:
            try:
                pdf_path = self._generate_pdf(report_data)
            except Exception as e:
                print(f"PDF generation error: {e}")

        return {
            "report_generated": True,
            "case_id": case_id,
            "files": {
                "docx": docx_path,
                "json": json_path,
                "pdf": pdf_path,
            },
            "report_data": report_data,
        }

    def _generate_docx(self, data: dict) -> str:
        doc = Document()

        # Title
        title = doc.add_heading("SHIELD – Credit Appraisal Memo (CAM)", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Case ID: {data['case_id']}")
        doc.add_paragraph(f"Generated: {data['generated_at']}")
        doc.add_paragraph("")

        # Borrower Info
        doc.add_heading("1. Borrower Information", level=2)
        doc.add_paragraph(f"Company Name: {data['company']}")
        doc.add_paragraph(f"Loan Amount Requested: ₹{data['loan_amount_cr']} Cr")

        # Financial Ratios
        doc.add_heading("2. Financial Ratios", level=2)
        ratios = data["financial_ratios"]
        doc.add_paragraph(f"Debt to Equity Ratio: {ratios.get('debt_to_equity', 'N/A')}")
        doc.add_paragraph(f"Current Ratio: {ratios.get('current_ratio', 'N/A')}")
        doc.add_paragraph(f"Profit Margin (EBITDA): {ratios.get('profit_margin', 'N/A')}")
        doc.add_paragraph(f"DSCR: {ratios.get('dscr', 'N/A')}x")

        # Risk Assessment
        doc.add_heading("3. Risk Assessment", level=2)
        risk = data["risk_assessment"]
        doc.add_paragraph(f"SHIELD Score: {risk.get('shield_score', 'N/A')}/100")
        doc.add_paragraph(f"Risk Level: {risk.get('risk_level', 'N/A')}")
        doc.add_paragraph(f"Default Probability: {risk.get('default_probability', 'N/A')}")

        # Breakdown
        breakdown = risk.get("breakdown", {})
        if breakdown:
            doc.add_heading("Score Breakdown", level=3)
            for key, val in breakdown.items():
                doc.add_paragraph(f"  {key.replace('_', ' ').title()}: {val}/100")

        # Fraud
        doc.add_heading("4. Fraud & Anomaly Analysis", level=2)
        fraud = data.get("fraud_analysis", {})
        doc.add_paragraph(f"Anomaly Detected: {fraud.get('anomaly_detected', False)}")
        flags = fraud.get("flags", [])
        if flags:
            doc.add_paragraph("Flags:")
            for f in flags:
                doc.add_paragraph(f"  • {f}")

        # Circular Trading
        doc.add_heading("5. Circular Trading Analysis", level=2)
        circ = data.get("circular_trading", {})
        doc.add_paragraph(f"Circular Trading Detected: {circ.get('fraud_detected', False)}")
        doc.add_paragraph(f"Confirmed Loops: {circ.get('cycle_count', 0)}")
        cycles = circ.get("cycles", [])
        if cycles:
            for i, cycle in enumerate(cycles):
                doc.add_paragraph(f"  Loop {i+1}: {' → '.join(cycle)}")

        # Decision
        doc.add_heading("6. Loan Decision", level=2)
        decision = data["loan_decision"]
        doc.add_paragraph(f"Decision: {decision.get('decision', 'N/A')}")
        reasons = decision.get("reasons", [])
        if reasons:
            doc.add_paragraph("Reasons:")
            for r in reasons:
                doc.add_paragraph(f"  • {r}")
        conditions = decision.get("conditions", [])
        if conditions:
            doc.add_paragraph("Pre-Disbursement Conditions:")
            for c in conditions:
                doc.add_paragraph(f"  • {c}")

        file_path = os.path.join(self.output_dir, f"{data['case_id']}_CAM.docx")
        doc.save(file_path)
        return file_path

    def _generate_json(self, data: dict) -> str:
        file_path = os.path.join(self.output_dir, f"{data['case_id']}_CAM.json")
        with open(file_path, "w") as f:
            json.dump(data, f, indent=2)
        return file_path

    def _generate_pdf(self, data: dict) -> str:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, _safe("SHIELD - Credit Appraisal Memo (CAM)"), ln=True, align="C")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, _safe(f"Case ID: {data['case_id']}"), ln=True)
        pdf.cell(0, 8, _safe(f"Company: {data['company']}"), ln=True)
        pdf.cell(0, 8, _safe(f"Loan Amount: Rs.{data['loan_amount_cr']} Cr"), ln=True)
        pdf.ln(4)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Financial Ratios", ln=True)
        pdf.set_font("Helvetica", "", 10)
        ratios = data["financial_ratios"]
        for k, v in ratios.items():
            pdf.cell(0, 7, _safe(f"  {k.replace('_', ' ').title()}: {v}"), ln=True)
        pdf.ln(3)

        risk = data["risk_assessment"]
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Risk Assessment", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, _safe(f"  SHIELD Score: {risk.get('shield_score')}/100"), ln=True)
        pdf.cell(0, 7, _safe(f"  Risk Level: {risk.get('risk_level')}"), ln=True)
        pdf.cell(0, 7, _safe(f"  Default Probability: {risk.get('default_probability')}"), ln=True)
        pdf.ln(3)

        fraud = data.get("fraud_analysis", {})
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Fraud Analysis", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, _safe(f"  Anomaly Detected: {fraud.get('anomaly_detected', False)}"), ln=True)
        for flag in fraud.get("flags", []):
            pdf.cell(0, 7, _safe(f"  * {flag}"), ln=True)
        pdf.ln(3)

        circ = data.get("circular_trading", {})
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Circular Trading", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, _safe(f"  Fraud Detected: {circ.get('fraud_detected', False)}"), ln=True)
        pdf.cell(0, 7, _safe(f"  Confirmed Loops: {circ.get('cycle_count', 0)}"), ln=True)
        pdf.ln(3)

        decision = data["loan_decision"]
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Loan Decision", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, _safe(f"  Decision: {decision.get('decision')}"), ln=True)
        for r in decision.get("reasons", []):
            pdf.cell(0, 7, _safe(f"  - {r}"), ln=True)
        for c in decision.get("conditions", []):
            pdf.cell(0, 7, _safe(f"  Condition: {c}"), ln=True)

        file_path = os.path.join(self.output_dir, f"{data['case_id']}_CAM.pdf")
        pdf.output(file_path)
        return file_path

