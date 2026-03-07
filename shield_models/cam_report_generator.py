from docx import Document


class CAMReportGenerator:

    def generate(self, company_name, ratios, risk_result, loan_decision):

        doc = Document()

        doc.add_heading("Credit Appraisal Memo (CAM)", level=1)

        doc.add_heading("Borrower Information", level=2)
        doc.add_paragraph(f"Company Name: {company_name}")

        doc.add_heading("Financial Ratios", level=2)
        doc.add_paragraph(f"Debt to Equity: {ratios['debt_to_equity']}")
        doc.add_paragraph(f"Current Ratio: {ratios['current_ratio']}")
        doc.add_paragraph(f"Profit Margin: {ratios['profit_margin']}")

        doc.add_heading("Risk Assessment", level=2)
        doc.add_paragraph(f"SHIELD Score: {risk_result['shield_score']}")
        doc.add_paragraph(f"Risk Level: {risk_result['risk_level']}")

        doc.add_heading("Loan Decision", level=2)
        doc.add_paragraph(f"Decision: {loan_decision['decision']}")

        if loan_decision["reasons"]:
            doc.add_paragraph("Reasons:")
            for r in loan_decision["reasons"]:
                doc.add_paragraph(f"- {r}")

        file_name = "CAM_Report.docx"

        doc.save(file_name)

        return {"report_generated": True, "file": file_name}